import re
import json
from xml.dom.minidom import Document
# Remove heavy imports from module level - will import when needed
import requests
import pdfplumber
from docx import Document
from lxml import html
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import os
from dotenv import load_dotenv
from .pii_anonymizer import PIIAnonymizer, anonymize_resume_text, anonymize_resume_data

# Load environment variables
load_dotenv()

# Get API key from environment
API_KEY = os.getenv("DEEPSEEK_API_KEY")

# Remove model loading from module level - will load when needed


def analysisPrompt(resume_text: str, job_details: str) -> list:
    """
    Creates a prompt for comparing resume against job posting.
    
    Args:
        resume_text: Processed resume text
        job_details: Processed job posting details
    
    Returns:
        list: Structured prompt for DeepSeek API with system and user messages
    """
    prompt = [
        {
            "role": "system",
            "content": "You are a recruiter that is comparing an applicant's resume to the job posting that you are hiring for. You will provide feedback in JSON format.",
        },
        {
            "role": "user",
            "content": f"""
                Here is the analysis of the applicant's resume: {resume_text}
                Here is the analysis of the job posting: {job_details}
                Now, run a comparison between the two analyses to provide feedback to the applicant.
                Create a JSON response with the following structure:
                {{
                    "strengths": [list of strengths compared to job posting],
                    "weaknesses": [list of weaknesses compared to job posting],
                    "improvement_tips": [list of tips to improve resume],
                    "keywords_missing": [list of keywords from job posting not found in resume],
                    "keywords_found": [list of keywords from job posting found in resume],
                    "match_score": number between 0 and 100
                }}
                    """,
        },
    ]
    return prompt


def resumeProcessorPrompt(resume_text: str) -> list:
    """
    Creates a structured prompt to extract resume information.
    
    Args:
        resume_text: Raw text extracted from resume
    
    Returns:
        list: Prompt template for extracting structured data from resume
    """
    prompt = [
        {
            "role": "system",
            "content": "You are a job applicant looking to see what information is in your resume.",
        },
        {
            "role": "user",
            "content": f"""Here is the text from your resume: {resume_text}

                    Look through the text that is given to find the following details and format them as a JSON object:
                    
                    {{
                        "name": "",
                        "contact_info": {{
                            "email": "",
                            "phone": "",
                            "zipCode": ""
                        }},
                        "work_experience": {{
                            "company": "",
                            "jobTitle": "",
                            "startDate": "",
                            "endDate": "",
                            "jobDescription": "",
                            "yearsOfExperience": ""
                        }},
                        "education": {{
                            "institution": "",
                            "highestDegree": "",
                            "fieldOfStudy": "",
                            "graduationYear": ""
                        }},
                        "projects": {{
                            "projectName": "",
                            "projectDescription": "",
                            "technologiesUsed": ""
                        }},
                        "certifications": {{
                            "certificationName": "",
                            "issuingOrganization": "",
                            "issueDate": "",
                            "expirationDate": ""
                        }},
                        "languages": {{
                            "language": "",
                            "proficiency": ""
                        }},
                        "linkedinUrl": "",
                        "skills": []
                    }}

                    For each field, include If a field has no value, set it to null. For skills, provide an array of strings.
                    """,
        },
    ]
    return prompt


def jobProcessorPrompt(job_posting: str) -> list:
    """
    Creates a prompt to extract job posting details.
    
    Args:
        job_posting: Raw text from job posting
    
    Returns:
        list: Prompt template for extracting job details
    """
    prompt = [
        {
            "role": "system",
            "content": "You are a job applicant seeking information from a job posting.",
        },
        {
            "role": "user",
            "content": f"""
                        Look through the text that is given to find the following details and include as much information as possible:
                        
                        title = string
                        description = string
                        qualifications = list[string]
                        skills = list[string]
                        responsibilities =list[string]
                        salary_range = string
                        location = string
                        posted_date = string
                        company_name = string
                        Here is the given text: {job_posting}

                        Once you have found the details, please put them into a JSON object. If nothing can be found for a field, set it to null.
                        """,
        },
    ]
    return prompt


def resumeJobDescAnalysis(resume_file_path: str, job_posting_url: str, anonymize_pii: bool = True) -> dict:
    """
    Analyzes resume against job posting and provides comparison with PII protection.
    
    Args:
        resume_file_path: Path to uploaded resume file
        job_posting_url: URL of job posting
        anonymize_pii: Whether to anonymize PII before sending to external AI
    
    Returns:
        dict: Analysis results including match score, strengths, weaknesses
    Raises:
        ValueError: If resume processing or job analysis fails
    """
    try:
        # Process the resume with anonymization option
        resume_data = processResume(resume_file_path, anonymize_pii=anonymize_pii)
        if not resume_data:
            raise ValueError("Failed to process resume")

        # Extract job description details (no PII in job postings)
        job_data = analyzeJobPosting(job_posting_url)
        if not job_data:
            raise ValueError("Failed to analyze job posting")

        # Prepare data for analysis
        resume_data_str = str(resume_data)
        job_data_str = str(job_data)
        
        # If we anonymized the resume, we need to anonymize the combined analysis data too
        pii_mapping = {}
        if anonymize_pii:
            print("🔒 Anonymizing combined analysis data...")
            anonymizer = PIIAnonymizer()
            
            # Anonymize the resume data string for analysis
            anonymized_resume_str, resume_mapping = anonymizer.anonymize_text(resume_data_str)
            pii_mapping.update(resume_mapping)
            
            resume_data_str = anonymized_resume_str

        # Implement job description and resume comparison here
        url = "https://api.deepseek.com/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {API_KEY}",
        }

        analysis_prompt = analysisPrompt(resume_data_str, job_data_str)
        
        if anonymize_pii:
            print("🚀 Sending anonymized analysis data to DeepSeek API...")
        else:
            print("⚠️  Sending original analysis data to DeepSeek API...")

        data = {
            "model": "deepseek-chat",
            "messages": analysis_prompt,
            "response_format": {"type": "json_object"},
            "stream": False,
        }
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            result = response.json()
            print("Raw API response:", json.dumps(result, indent=2))
            try:
                if (
                    "choices" in result
                    and len(result["choices"]) > 0
                    and "message" in result["choices"][0]
                    and "content" in result["choices"][0]["message"]
                ):
                    content = json.loads(result["choices"][0]["message"]["content"])
                    
                    # If we anonymized, deanonymize the response
                    if anonymize_pii and pii_mapping:
                        print("🔓 Deanonymizing analysis response...")
                        anonymizer = PIIAnonymizer()
                        
                        # Deanonymize all string fields in the response
                        for key, value in content.items():
                            if isinstance(value, str):
                                content[key] = anonymizer.deanonymize_text(value, pii_mapping)
                            elif isinstance(value, list):
                                content[key] = [
                                    anonymizer.deanonymize_text(item, pii_mapping) if isinstance(item, str) else item
                                    for item in value
                                ]
                        
                        # Add metadata
                        content["_pii_anonymized"] = True
                        content["_analysis_with_privacy_protection"] = True
                    else:
                        content["_pii_anonymized"] = False
                        content["_analysis_with_privacy_protection"] = False
                    
                    print("Parsed content:", json.dumps(content, indent=2))
                    return content
                else:
                    raise KeyError("Missing expected keys in API response")
            except KeyError as e:
                print(f"Error processing API response: {str(e)}")
                choices = result.get("choices", [])
                if (
                    choices
                    and "message" in choices[0]
                    and "content" in choices[0]["message"]
                ):
                    content = json.loads(choices[0]["message"]["content"])
                    print(
                        "Parsed content (from error handler):",
                        json.dumps(content, indent=2),
                    )
                    return content
                else:
                    print(
                        "Unexpected response format: Missing 'choices' or nested keys."
                    )
            except Exception as e:
                print(f"Error processing API response: {str(e)}")
                print("Response content:", result["choices"][0]["message"]["content"])
            return json.loads(result["choices"][0]["message"]["content"])
        else:
            print("Request failed, error code:", response.status_code)
            print("Response content:", response.text)

    except Exception as e:
        print(f"Error in resumeJobDescAnalysis: {str(e)}")
        return {
            "error": str(e),
            "status": "failed",
            "resume_processed": bool(resume_data) if 'resume_data' in locals() else False,
            "job_data_processed": bool(job_data) if 'job_data' in locals() else False,
        }


def processResumeFromContent(file_content: bytes, filename: str, anonymize_pii: bool = True) -> dict:
    """
    Extracts and processes text from resume file content (in-memory processing).
    
    Args:
        file_content: Raw file content as bytes
        filename: Original filename to determine file type
        anonymize_pii: Whether to anonymize PII before sending to external AI
    
    Returns:
        dict: Structured resume data with anonymization metadata if applicable
    Raises:
        ValueError: If file format is not supported
    """
    import io
    import tempfile
    
    resume_text = ""
    
    # Determine file type from filename
    if filename.lower().endswith(".pdf"):
        # Process PDF from memory
        pdf_file = io.BytesIO(file_content)
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    resume_text += page_text
                    
    elif filename.lower().endswith(".docx"):
        # Process DOCX from memory  
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        resume_text = "\n".join([p.text for p in doc.paragraphs])
        
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX files are supported.")
    
    if not resume_text.strip():
        raise ValueError("No text could be extracted from the resume file.")
    
    # Clean up text
    resume_text = re.sub(r"\s+", " ", resume_text).strip()
    
    # Anonymize PII if requested
    pii_mapping = {}
    processed_text = resume_text
    
    if anonymize_pii:
        print("🔒 Anonymizing PII before sending to external AI service...")
        anonymizer = PIIAnonymizer()
        processed_text, pii_mapping = anonymizer.anonymize_text(resume_text)
        
        # Create anonymization report
        anonymization_report = anonymizer.create_anonymization_report(pii_mapping)
        print(f"📊 Anonymization Report: {anonymization_report['total_items']} PII items anonymized")
        print(f"   Types: {anonymization_report['types']}")

    # Process with DeepSeek API
    url = "https://api.deepseek.com/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"}

    prompt = resumeProcessorPrompt(processed_text)
    
    if anonymize_pii:
        print("🚀 Sending anonymized resume to DeepSeek API...")
    else:
        print("⚠️  Sending original resume text to DeepSeek API...")
    
    data = {
        "model": "deepseek-chat",
        "messages": prompt,
        "response_format": {"type": "json_object"},
        "stream": False,
    }
    
    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            result = response.json()
            print("Raw API response:", json.dumps(result, indent=2))
            
            if (
                "choices" in result
                and len(result["choices"]) > 0
                and "message" in result["choices"][0]
                and "content" in result["choices"][0]["message"]
            ):
                content = json.loads(result["choices"][0]["message"]["content"])
                
                # If we anonymized, we need to deanonymize the response
                if anonymize_pii and pii_mapping:
                    print("🔓 Deanonymizing AI response...")
                    anonymizer = PIIAnonymizer()
                    content = anonymizer.deanonymize_data(content, pii_mapping)
                
                # Add processing metadata
                content['pii_anonymized'] = anonymize_pii
                content['original_text_length'] = len(resume_text)
                if anonymize_pii:
                    content['anonymized_text_length'] = len(processed_text)
                    content['pii_items_anonymized'] = anonymization_report['total_items']
                
                return content
            else:
                raise ValueError("Invalid response format from DeepSeek API")
        else:
            raise ValueError(f"DeepSeek API error: {response.status_code} - {response.text}")
    except Exception as e:
        print(f"Error processing resume: {str(e)}")
        # Return basic structure on error
        return {
            "error": str(e),
            "pii_anonymized": anonymize_pii,
            "original_text_length": len(resume_text),
            "processing_failed": True
        }

def processResume(resume_file_path: str, anonymize_pii: bool = True) -> dict:
    """
    Extracts and processes text from resume files with optional PII anonymization.
    
    Args:
        resume_file_path: Path to PDF or DOCX resume
        anonymize_pii: Whether to anonymize PII before sending to external AI
    
    Returns:
        dict: Structured resume data with anonymization metadata if applicable
    Raises:
        ValueError: If file format is not supported
    """
    resume_text = ""
    if resume_file_path.endswith(".pdf"):
        with pdfplumber.open(f"{resume_file_path}") as pdf:
            for page in pdf.pages:
                resume_text += page.extract_text()
    elif resume_file_path.endswith(".docx"):
        doc = Document(resume_file_path)
        resume_text = "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Use PDF or DOCX.")

    resume_text = re.sub(r"\s+", " ", resume_text).strip()
    
    # Anonymize PII if requested
    pii_mapping = {}
    processed_text = resume_text
    
    if anonymize_pii:
        print("🔒 Anonymizing PII before sending to external AI service...")
        anonymizer = PIIAnonymizer()
        processed_text, pii_mapping = anonymizer.anonymize_text(resume_text)
        
        # Create anonymization report
        anonymization_report = anonymizer.create_anonymization_report(pii_mapping)
        print(f"📊 Anonymization Report: {anonymization_report['total_items']} PII items anonymized")
        print(f"   Types: {anonymization_report['types']}")

    url = "https://api.deepseek.com/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"}

    prompt = resumeProcessorPrompt(processed_text)
    
    if anonymize_pii:
        print("🚀 Sending anonymized resume to DeepSeek API...")
    else:
        print("⚠️  Sending original resume text to DeepSeek API...")
    
    data = {
        "model": "deepseek-chat",
        "messages": prompt,
        "response_format": {"type": "json_object"},
        "stream": False,
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        result = response.json()
        print("Raw API response:", json.dumps(result, indent=2))
        try:
            if (
                "choices" in result
                and len(result["choices"]) > 0
                and "message" in result["choices"][0]
                and "content" in result["choices"][0]["message"]
            ):
                content = json.loads(result["choices"][0]["message"]["content"])
                
                # If we anonymized, we need to deanonymize the response
                if anonymize_pii and pii_mapping:
                    print("🔓 Deanonymizing AI response...")
                    anonymizer = PIIAnonymizer()
                    
                    # Convert content back to string, deanonymize, then parse back
                    content_str = json.dumps(content)
                    deanonymized_str = anonymizer.deanonymize_text(content_str, pii_mapping)
                    
                    try:
                        content = json.loads(deanonymized_str)
                    except json.JSONDecodeError:
                        print("Warning: Could not parse deanonymized content as JSON, using original")
                    
                    # Add metadata about anonymization
                    content["_pii_anonymized"] = True
                    content["_anonymization_report"] = anonymization_report
                else:
                    content["_pii_anonymized"] = False
                
                print("Parsed content:", json.dumps(content, indent=2))
                return content
            else:
                raise KeyError("Missing expected keys in API response")
        except KeyError as e:
            print(f"Error processing API response: {str(e)}")
            choices = result.get("choices", [])
            if (
                choices
                and "message" in choices[0]
                and "content" in choices[0]["message"]
            ):
                content = json.loads(choices[0]["message"]["content"])
                print(
                    "Parsed content (from error handler):",
                    json.dumps(content, indent=2),
                )
                return content
            else:
                print("Unexpected response format: Missing 'choices' or nested keys.")
        except Exception as e:
            print(f"Error processing API response: {str(e)}")
            print("Response content:", result["choices"][0]["message"]["content"])
        return json.loads(result["choices"][0]["message"]["content"])
    else:
        print("Request failed, error code:", response.status_code)
        print("Response content:", response.text)



def analyzeJobPosting(job_posting: str) -> dict:
    """
    Analyzes job posting text using DeepSeek API.
    
    Args:
        job_posting: Raw job posting text
    
    Returns:
        dict: Structured job posting data
    """
    url = "https://api.deepseek.com/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"}

    prompt = jobProcessorPrompt(job_posting)
    print("Sending prompt to API:", json.dumps(prompt, indent=2))

    data = {
        "model": "deepseek-chat",  # Use 'deepseek-reasoner' for R1 model or 'deepseek-chat' for V3 model
        "messages": prompt,
        "response_format": {"type": "json_object"},
        "stream": False,  # Disable streaming
    }

    response = requests.post(url, headers=headers, json=data)

    if response.status_code == 200:
        result = response.json()
        print("Raw API response:", json.dumps(result, indent=2))
        try:
            if (
                "choices" in result
                and len(result["choices"]) > 0
                and "message" in result["choices"][0]
                and "content" in result["choices"][0]["message"]
            ):
                content = json.loads(result["choices"][0]["message"]["content"])
                print("Parsed content:", json.dumps(content, indent=2))
                return content
            else:
                raise KeyError("Missing expected keys in API response")
        except KeyError as e:
            print(f"Error processing API response: {str(e)}")
            choices = result.get("choices", [])
            if (
                choices
                and "message" in choices[0]
                and "content" in choices[0]["message"]
            ):
                content = json.loads(choices[0]["message"]["content"])
                print(
                    "Parsed content (from error handler):",
                    json.dumps(content, indent=2),
                )
                return content
            else:
                print("Unexpected response format: Missing 'choices' or nested keys.")
        except Exception as e:
            print(f"Error processing API response: {str(e)}")
            print("Response content:", result["choices"][0]["message"]["content"])
        return json.loads(result["choices"][0]["message"]["content"])
    else:
        print(
            f"Error: analyzeJobPosting\nRequest failed, error code:{response.status_code}\n"
        )


def getRawText(tree: html.HtmlElement) -> str:
    """
    Extracts and cleans text from HTML.
    
    Args:
        tree: Parsed HTML element tree
    
    Returns:
        str: Cleaned text content
    """
    # Remove scripts/styles (XPath)
    for tag in tree.xpath("//script | //style | //noscript | //svg"):
        tag.getparent().remove(tag)
    raw_text = tree.text_content()  # Extracts all text, including whitespace
    text = re.sub(r"\s+", " ", raw_text).strip()  # Normalize whitespace
    return text


def ner(text: str) -> dict:
    """
    Named Entity Recognition (NER) using the DSLIM BERT model

    Args:
        text (str): Input text for NER

    Returns:
        dict: NER results
    """
    from transformers import (
        AutoTokenizer,
        AutoModelForTokenClassification,
        pipeline,
    )
    tokenizer = AutoTokenizer.from_pretrained("dslim/bert-base-NER")
    model = AutoModelForTokenClassification.from_pretrained("dslim/bert-base-NER")
    ner_pipeline = pipeline("ner", model=model, tokenizer=tokenizer)

    return ner_pipeline(text)


# Ask a question and get an answer
def askQuestion(text: str, question: str) -> str:
    """
    Uses QA model to extract specific information.
    
    Args:
        text: Context text
        question: Question to answer
    
    Returns:
        str: Extracted answer from text
    """
    from transformers import (
        AutoTokenizer,
        AutoModelForTokenClassification,
        pipeline,
    )
    qa_pipeline = pipeline(
        "question-answering", model="distilbert-base-cased-distilled-squad"
    )
    answer = qa_pipeline(question=question, context=text)
    return answer["answer"]


# Extract qualifications, responsibilities, and salary range from job description
def extractQAFields(text: str) -> dict:
    """
    Extracts specific fields from job description.
    
    Args:
        text: Job description text
    
    Returns:
        dict: Extracted fields (description, qualifications, skills, etc.)
    """
    questions = {
        # Question to extract the overall job description
        "description": "What is the job description?",
        # Question to identify the qualifications required for the job
        "qualifications": "What are the required qualifications from the job description?",
        # Question to list the skills needed for the job
        "skills": "What are the required skills from the job description?",
        # Question to determine the key responsibilities of the job
        "responsibilities": "What are the key responsibilities from the job description?",
        # Question to extract the salary range offered for the job
        "salary": "What is the salary range?",
    }
    return {k: askQuestion(text, q) for k, q in questions.items()}

# Extract job description from a URL using Selenium
def extractJobDescription(url: str) -> dict:
    """
    Scrapes job posting content using Selenium.
    
    Args:
        url: Job posting URL
    
    Returns:
        dict: Processed job posting data or error details
    """
    import uuid
    import tempfile
    
    # Validate API key early
    if not API_KEY:
        print("❌ DEEPSEEK_API_KEY environment variable not set")
        return {
            "url": url,
            "status_code": 500,
            "description": "DeepSeek API key not configured",
            "full_text": "",
            "html": "",
        }
    
    driver = None
    temp_user_data_dir = None
    
    try:
        # Create unique temporary directory for this session
        temp_user_data_dir = tempfile.mkdtemp(prefix=f"chrome_user_data_{uuid.uuid4().hex[:8]}_")
        
        # Set up Chrome options with unique user data directory
        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--no-sandbox") 
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--disable-software-rasterizer")
        chrome_options.add_argument("--disable-background-timer-throttling")
        chrome_options.add_argument("--disable-backgrounding-occluded-windows")
        chrome_options.add_argument("--disable-renderer-backgrounding")
        chrome_options.add_argument("--disable-features=TranslateUI")
        chrome_options.add_argument("--disable-ipc-flooding-protection")
        chrome_options.add_argument(f"--user-data-dir={temp_user_data_dir}")
        chrome_options.add_argument("--single-process")
        chrome_options.add_argument("--disable-extensions")
        chrome_options.add_argument("--disable-plugins")
        chrome_options.add_argument("--disable-images")
        chrome_options.add_argument("--disable-javascript")
        chrome_options.add_argument("--disable-web-security")
        chrome_options.binary_location = "/usr/bin/chromium"

        # Initialize the driver with explicit service
        service = Service(executable_path="/usr/bin/chromedriver")
        driver = webdriver.Chrome(service=service, options=chrome_options)

        # Load the page and wait for content to load
        driver.get(url)

        # Wait for the page to load by checking for the presence of a specific element
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.TAG_NAME, "body"))
        )

        # Get the page source after JavaScript execution
        job_posting_html = driver.page_source
        
        # Parse HTML and extract text
        tree = html.fromstring(job_posting_html)
        job_posting = getRawText(tree)

        # Process with your existing analysis
        job_details_deepseek = analyzeJobPosting(job_posting)
        if job_details_deepseek:
            return job_details_deepseek
        else:
            return {
                "url": url,
                "status_code": 200,
                "description": "No job details extracted",
                "full_text": job_posting,
                "html": job_posting_html,
            }

    except Exception as e:
        print(f"WebDriver error for {url}: {str(e)}")
        return {
            "url": url,
            "status_code": 500,
            "description": f"Error extracting job posting: {str(e)}",
            "full_text": "",
            "html": "",
        }
    
    finally:
        # Always cleanup resources
        if driver:
            try:
                driver.quit()
                print("✅ WebDriver session closed successfully")
            except Exception as e:
                print(f"⚠️  Error closing WebDriver: {e}")
        
        # Clean up temporary user data directory
        if temp_user_data_dir:
            try:
                import shutil
                shutil.rmtree(temp_user_data_dir, ignore_errors=True)
                print(f"✅ Cleaned up temporary directory: {temp_user_data_dir}")
            except Exception as e:
                print(f"⚠️  Error cleaning temp directory {temp_user_data_dir}: {e}")
